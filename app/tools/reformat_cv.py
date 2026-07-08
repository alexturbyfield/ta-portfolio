from __future__ import annotations

import re
import tempfile
import zipfile
from xml.etree import ElementTree as ET
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\aturb\ta-portfolio\app")
SOURCE = ROOT / "Alex_Turbyfield_CV_original.docx"
OUTPUT = ROOT / "Alex_Turbyfield_CV_revised.docx"
FINAL_OUTPUT = ROOT / "Alex_Turbyfield_CV_academic.docx"


SECTION_TITLES = [
    "Education",
    "Awards/Events/Committees/Leadership",
    "Teaching Experience",
    "Guest Lectures/Workshops",
    "Professional Software and Visual Effects Experience",
    "Animation Experience",
    "Illustration and Graphic Design Experience",
    "Puppeteer Experience",
    "Acting Experience",
    "Theatre and Film Directing Experience",
    "Other Theatre and Film Experience",
    "Press and Media Appearances",
]


SECTION_ORDER = [
    "Education",
    "Teaching Areas",
    "Teaching Experience",
    "Guest Lectures/Workshops",
    "Professional Software and Visual Effects Experience",
    "Awards/Events/Committees/Leadership",
    "Animation Experience",
    "Illustration and Graphic Design Experience",
    "Puppeteer Experience",
    "Acting Experience",
    "Theatre and Film Directing Experience",
    "Other Theatre and Film Experience",
    "Press and Media Appearances",
]


SECTION_LABELS = {
    "Awards/Events/Committees/Leadership": "Honors, Awards, and Leadership",
    "Guest Lectures/Workshops": "Invited Lectures and Workshops",
    "Professional Software and Visual Effects Experience": "Professional Appointments and Industry Experience",
    "Animation Experience": "Animation and Interactive Media Projects",
    "Illustration and Graphic Design Experience": "Illustration and Graphic Design",
    "Puppeteer Experience": "Puppetry and Performance Experience",
    "Acting Experience": "Acting and Voice Performance",
    "Theatre and Film Directing Experience": "Theatre and Film Directing",
    "Other Theatre and Film Experience": "Additional Theatre and Film Experience",
    "Press and Media Appearances": "Press and Media Coverage",
}


TEACHING_AREAS = [
    "Animation fundamentals and production",
    "Technical art for games and interactive media",
    "3D modeling, rigging, and visual development",
    "Unity game development and interactive prototyping",
    "Autodesk Maya workflows for artists and animators",
    "Digital media for non-computer science majors",
    "Sound and media design for theatre, film, and animation",
]


ORG_HEADINGS = {
    "Teaching Experience": {"University of Georgia"},
    "Guest Lectures/Workshops": {
        "University of West Georgia",
        "Pokémon World Championships",
        "re:imagineATL",
        "University of Georgia",
        "Georgia Game Developers Association",
    },
    "Animation Experience": {"University of Georgia"},
    "Press and Media Appearances": {"SYSTEMS"},
}


def add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    new_run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def normalize_whitespace(text: str) -> str:
    text = text.replace("\u200b", "")
    text = text.replace("\t", " ")
    text = text.replace("–", "—")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def clean_text(text: str) -> str:
    text = normalize_whitespace(text)

    replacements = {
        "3d": "3D",
        "Maya 3d": "Maya 3D",
        "animation..": "animation.",
        "Sound Design in Theatre, Television, Film, and Animation": "sound design in theatre, television, film, and animation",
        "nParticles in Autodesk Maya 3d": "nParticles in Autodesk Maya 3D",
        "Software used in include": "Software used include",
        "based off of": "based on",
        "shader programing": "shader programming",
        "graphics,animations, 3D models,,": "graphics, animations, 3D models,",
        "Mac OS and IOS": "macOS and iOS",
        "tv pilot": "TV pilot",
        "Award Winning": "award-winning",
        "What’s Your Favorite?": "\"What's Your Favorite?\"",
        "The Art of Pokémon Go": "The Art of Pokémon GO",
        "Pokémon Go": "Pokémon GO",
        "Digital Media Supervisor of International Animation Society (Association Internationale du Film d'Animation)- South Chapter affiliated with UNESCO.": "Digital Media Supervisor for ASIFA-South, the International Animated Film Association's South Chapter, affiliated with UNESCO.",
        "Developer for the ASIFA-SOUTH Animation Conference and Festival, including augmented reality application development, coding, and management of animation and film implementation.": "Developer for the ASIFA-South Animation Conference and Festival, including augmented reality application development, programming, and implementation management for animation and film projects.",
        "Handles digital media for special events.": "Managed digital media for special events.",
        "With Developer Stephen Borden,": "With developer Stephen Borden,",
        "3D modelling": "3D modeling",
        "University of West Georgia – 2007": "University of West Georgia—2007",
        "University of West Georgia — 2007": "University of West Georgia—2007",
        "Georgia Forward Young Gamechangers - Fall 2016 Class": "Georgia Forward Young Gamechangers—Fall 2016 Class",
        "Spring 2008": "Spring 2008",
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    text = re.sub(r"\b11-18\b", "11–18", text)
    text = re.sub(r"\b2010-2012\b", "2010–2012", text)
    text = re.sub(r"\b2007-2008\b", "2007–2008", text)
    text = re.sub(r"\b2006-2007\b", "2006–2007", text)
    text = re.sub(r"\b2009-2010\b", "2009–2010", text)
    text = re.sub(r"\b2003-2004\b", "2003–2004", text)
    text = re.sub(r"\b2017-2020\b", "2017–2020", text)
    text = re.sub(r"\b2017-present\b", "2017–present", text)
    text = re.sub(r"\b2014-present\b", "2014–present", text)
    text = re.sub(r"\b2013-September 2014\b", "August 2013–September 2014", text)
    text = text.replace("—August August 2013–September 2014", "—August 2013–September 2014")
    text = re.sub(r"—Fall  +", "—Fall ", text)
    text = text.replace("The Grapes of Wrath(", "The Grapes of Wrath (")
    text = text.replace("basics of  ", "basics of ")

    return text.strip()


def load_contact_info(doc: Document) -> tuple[str, list[str]]:
    table = doc.tables[0]
    name = normalize_whitespace(table.cell(0, 0).text)
    raw_contact = [normalize_whitespace(line) for line in table.cell(0, 1).text.splitlines() if normalize_whitespace(line)]

    city_state = ""
    phone = ""
    email = ""
    website = ""

    for line in raw_contact:
        if "@" in line:
            email = line
        elif line.startswith("www."):
            website = line.replace("www.", "")
        elif re.search(r"\(\d{3}\)", line):
            phone = line
        elif "Redmond" in line:
            city_state = line.replace(", 98052", "")

    contact_line = " | ".join([part for part in [city_state, phone, email, website] if part])
    return name, [contact_line]


def load_sections(doc: Document) -> dict[str, list[str]]:
    items: list[str] = []
    for para in doc.paragraphs:
        text = normalize_whitespace(para.text)
        if text:
            items.append(text)

    merged: list[str] = []
    for text in items:
        if text == "2008" and merged and merged[-1].endswith("—Summer"):
            merged[-1] = f"{merged[-1]} 2008"
            continue
        if text.startswith("Respect, Personal, Local, and Global") and merged and merged[-1].endswith("Episodes in Sexuality:"):
            merged[-1] = f"{merged[-1]} {text}"
            continue
        merged.append(text)

    sections: dict[str, list[str]] = {}
    current = None
    for text in merged:
        if text in SECTION_TITLES:
            current = text
            sections[current] = []
        elif text == "Curriculum Vitae":
            continue
        elif current:
            sections[current].append(clean_text(text))

    return sections


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)

    p_format = normal.paragraph_format
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(8)
    p_format.line_spacing = 1.15

    def ensure_paragraph_style(name: str) -> None:
        if name not in doc.styles:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    for name in ["SectionHeading", "OrgHeading", "EntryTitle", "EntryBody", "CompactBullet", "SmallMeta"]:
        ensure_paragraph_style(name)

    section_heading = doc.styles["SectionHeading"]
    section_heading.base_style = normal
    section_heading.font.name = "Arial"
    section_heading.font.size = Pt(12)
    section_heading.font.bold = True
    section_heading.paragraph_format.space_before = Pt(16)
    section_heading.paragraph_format.space_after = Pt(4)
    section_heading.paragraph_format.keep_with_next = True

    org_heading = doc.styles["OrgHeading"]
    org_heading.base_style = normal
    org_heading.font.name = "Arial"
    org_heading.font.size = Pt(11)
    org_heading.font.bold = True
    org_heading.paragraph_format.space_before = Pt(10)
    org_heading.paragraph_format.space_after = Pt(2)
    org_heading.paragraph_format.keep_with_next = True

    entry_title = doc.styles["EntryTitle"]
    entry_title.base_style = normal
    entry_title.font.name = "Arial"
    entry_title.font.size = Pt(11)
    entry_title.font.bold = True
    entry_title.paragraph_format.space_after = Pt(2)
    entry_title.paragraph_format.keep_with_next = True
    entry_title.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)

    entry_body = doc.styles["EntryBody"]
    entry_body.base_style = normal
    entry_body.font.name = "Arial"
    entry_body.font.size = Pt(11)
    entry_body.paragraph_format.left_indent = Inches(0.2)
    entry_body.paragraph_format.space_after = Pt(5)

    compact_bullet = doc.styles["CompactBullet"]
    compact_bullet.base_style = doc.styles["List Bullet"]
    compact_bullet.font.name = "Arial"
    compact_bullet.font.size = Pt(11)
    compact_bullet.paragraph_format.left_indent = Inches(0.25)
    compact_bullet.paragraph_format.space_after = Pt(3)
    compact_bullet.paragraph_format.line_spacing = 1.15

    small_meta = doc.styles["SmallMeta"]
    small_meta.base_style = normal
    small_meta.font.name = "Arial"
    small_meta.font.size = Pt(10)
    small_meta.font.italic = True
    small_meta.font.color.rgb = RGBColor(85, 85, 85)
    small_meta.paragraph_format.space_after = Pt(3)


def add_name_block(doc: Document, name: str, contact_lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(name)
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.bold = True

    for line in contact_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Curriculum Vitae")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_section_heading(doc: Document, title: str) -> None:
    label = SECTION_LABELS.get(title, title)
    p = doc.add_paragraph(style="SectionHeading")
    run = p.add_run(label.upper())
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = True


def build_teaching_areas(doc: Document) -> None:
    add_section_heading(doc, "Teaching Areas")
    add_bullet_items(doc, TEACHING_AREAS)


def add_bullet_items(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="CompactBullet")


def add_entry_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="EntryTitle")
    p.add_run(text)


def add_entry_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="EntryBody")
    p.add_run(text)


def add_meta_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="SmallMeta")
    p.add_run(text)


def build_education(doc: Document, items: list[str]) -> None:
    add_section_heading(doc, "Education")
    for item in items:
        if item.startswith(("M.F.A.", "B.A.")):
            add_entry_title(doc, item)
        else:
            add_entry_body(doc, item)


def build_teaching(doc: Document, items: list[str], section_name: str) -> None:
    add_section_heading(doc, section_name)
    orgs = ORG_HEADINGS[section_name]
    current_org = None
    idx = 0
    while idx < len(items):
        item = items[idx]
        if item in orgs:
            current_org = item
            p = doc.add_paragraph(style="OrgHeading")
            p.add_run(item)
            idx += 1
            continue
        add_entry_title(doc, item)
        if idx + 1 < len(items) and items[idx + 1] not in orgs:
            add_entry_body(doc, items[idx + 1])
            idx += 2
        else:
            idx += 1


def is_professional_entry(text: str) -> bool:
    if "." in text:
        return False
    return "—" in text and not text.startswith(("Project:", "March ", "Spring ", "Summer ", "Fall ", "Winter "))


def split_trailing_date(text: str) -> tuple[str, str] | None:
    match = re.search(
        r"^(.*?)(—(?:January|February|March|April|May|June|July|August|September|October|November|December|Spring|Summer|Fall|Winter)\s.+)$",
        text,
    )
    if not match or "." not in match.group(1):
        return None
    body = match.group(1).rstrip(" .")
    meta = match.group(2).lstrip("—").strip()
    return body + ".", meta


def build_professional(doc: Document, items: list[str]) -> None:
    add_section_heading(doc, "Professional Software and Visual Effects Experience")
    idx = 0
    while idx < len(items):
        item = items[idx]
        if is_professional_entry(item):
            add_entry_title(doc, item)
            idx += 1
            while idx < len(items) and not is_professional_entry(items[idx]):
                detail = items[idx]
                if detail.startswith(("Project:", "March ", "April ", "May ", "June ", "July ", "August ", "September ", "October ", "November ", "December ", "January ", "February ")) or re.fullmatch(r"[A-Z][a-z]+ \d{4}", detail):
                    add_meta_line(doc, detail)
                else:
                    split = split_trailing_date(detail)
                    if split:
                        body, meta = split
                        add_entry_body(doc, body)
                        add_meta_line(doc, meta)
                    else:
                        add_entry_body(doc, detail)
                idx += 1
        else:
            idx += 1


def build_list_section(doc: Document, section_name: str, items: list[str]) -> None:
    add_section_heading(doc, section_name)
    orgs = ORG_HEADINGS.get(section_name, set())
    buffer: list[str] = []
    for item in items:
        if item in orgs:
            if buffer:
                add_bullet_items(doc, buffer)
                buffer = []
            p = doc.add_paragraph(style="OrgHeading")
            p.add_run(item)
        else:
            buffer.append(item)
    if buffer:
        add_bullet_items(doc, buffer)


def split_press_item(item: str) -> tuple[str, str] | None:
    match = re.search(r"(https?://\S+)$", item)
    if not match:
        return None
    url = match.group(1)
    text = item[: match.start()].strip()
    return text, url


def build_press(doc: Document, items: list[str]) -> None:
    add_section_heading(doc, "Press and Media Appearances")
    current_org = None
    for item in items:
        if item in ORG_HEADINGS["Press and Media Appearances"]:
            current_org = item
            p = doc.add_paragraph(style="OrgHeading")
            p.add_run(item)
            continue
        p = doc.add_paragraph(style="CompactBullet")
        parsed = split_press_item(item)
        if parsed:
            text, url = parsed
            p.add_run(f"{text} ")
            add_hyperlink(p, url, url)
        else:
            p.add_run(item)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_text = OxmlElement("w:t")
    fld_text.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(fld_text)
    run._r.append(fld_end)


def configure_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Page ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(102, 102, 102)
    add_page_number(paragraph)
    for extra in paragraph.runs[1:]:
        extra.font.name = "Arial"
        extra.font.size = Pt(9)
        extra.font.color.rgb = RGBColor(102, 102, 102)


def strip_title_style_border(docx_path: Path) -> None:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        with zipfile.ZipFile(docx_path, "r") as src_zip:
            src_zip.extractall(tmpdir_path)

        styles_path = tmpdir_path / "word" / "styles.xml"
        tree = ET.parse(styles_path)
        root = tree.getroot()

        for style in root.findall("w:style", ns):
            if style.get(f"{{{ns['w']}}}styleId") != "Title":
                continue
            ppr = style.find("w:pPr", ns)
            if ppr is None:
                continue
            for border in list(ppr.findall("w:pBdr", ns)):
                ppr.remove(border)

        tree.write(styles_path, encoding="utf-8", xml_declaration=True)

        with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as dest_zip:
            for file_path in tmpdir_path.rglob("*"):
                if file_path.is_file():
                    dest_zip.write(file_path, file_path.relative_to(tmpdir_path))


def main() -> None:
    source = Document(str(SOURCE))
    name, contact_lines = load_contact_info(source)
    sections = load_sections(source)

    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    configure_footer(doc)
    add_name_block(doc, name, contact_lines)

    for section_name in SECTION_ORDER:
        if section_name == "Teaching Areas":
            build_teaching_areas(doc)
            continue
        items = sections.get(section_name, [])
        if not items:
            continue
        if section_name == "Education":
            build_education(doc, items)
        elif section_name in {"Teaching Experience", "Guest Lectures/Workshops"}:
            build_teaching(doc, items, section_name)
        elif section_name == "Professional Software and Visual Effects Experience":
            build_professional(doc, items)
        elif section_name == "Press and Media Appearances":
            build_press(doc, items)
        else:
            build_list_section(doc, section_name, items)

    if doc.paragraphs:
        last = doc.paragraphs[-1]
        last.paragraph_format.space_after = Pt(0)

    doc.save(str(OUTPUT))
    strip_title_style_border(OUTPUT)
    OUTPUT.replace(FINAL_OUTPUT)


if __name__ == "__main__":
    main()
